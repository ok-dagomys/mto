import docx
from docx.shared import Inches


doc = docx.Document('template.docx')


def delete_row_in_table(table, row):
    doc.tables[table]._tbl.remove(doc.tables[table].rows[row]._tr)


def add_table(dataframe):
    table = doc.add_table(dataframe.shape[0] + 1, dataframe.shape[1], style='Table Grid')

    # for j in range(dataframe.shape[-1]):
    #     table.cell(0, j).text = dataframe.columns[j]

    for i in range(dataframe.shape[0]):
        for j in range(dataframe.shape[-1]):
            table.cell(i + 1, j).text = str(dataframe.values[i, j])

    widths = (Inches(0.4), Inches(6), Inches(0.55), Inches(0.55))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    delete_row_in_table(1, 0)
    return table


def move_table_after(table, search_paragraph):
    tbl, p = table._tbl, search_paragraph._p
    p.addnext(tbl)


def delete_paragraph(del_paragraph):
    p = del_paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def create_task(df):
    for search in doc.paragraphs:
        if 'insert table here' in search.text:
            paragraph = search
            new_table = add_table(df)
            move_table_after(new_table, paragraph)
            delete_paragraph(paragraph)

    doc.save('result.docx')
